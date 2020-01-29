import requests

from bs4 import BeautifulSoup
from docx import Document


def set_docx (title, timestamp, intro, paragraphs, filename):
    '''
    Sets the appropriate style to set the input as an article and saves as docx
    '''    
    document = Document()

    document.add_heading(title, level=1)
    document.add_paragraph(timestamp, style="Quote")    
    document.add_heading(intro, level=3)
    document.add_paragraph('')
    
    for para in paragraphs:
        document.add_paragraph(' '*12 + para)
        
    document.save(filename + '.docx')


def get_article_TH (url, filename):
    '''
    Gets the article from the "The Hindu - Today's Paper" webpage specified by the url and saves as pdf by the name filename
    '''

    # Fetch page data from url
    soup = BeautifulSoup(requests.get(url).text, 'html.parser')

    # Fetch relevant contents
    try:
        article_title = soup.find_all("h1", class_="title")[0].text.strip()
        article_date = soup.find_all(
            "span", class_="ksl-time-stamp")[0].text.strip()[:-9]
        article_intro = soup.find_all("h2", class_="intro")[0].text.strip()
        article_body = soup.find_all("h2", class_="intro")[
            0].parent.find_all("div")[0].find_all("p")

        # Save article as docx
        set_docx(article_title, article_date, article_intro, [
        para.text for para in article_body], filename)
    except:
        pass

def get_article_IE (url, filename):
    '''
    Gets the article from the "Indian Express" Editorials webpage specified by the url and saves as pdf by the name filename
    '''

    # Fetch page date from url
    soup = BeautifulSoup(requests.get(url).text, 'html.parser')

    # Fetch relevant content
    article_title = soup.find_all("h1", class_="m-story-header__title")[0].text.strip()
    article_date = soup.find_all(class_="m-story-meta__credit")[0].text.strip()[24:-10]
    article_intro = soup.find_all("h2", class_="m-story-header__intro")[0].text.strip()
    article_body = soup.find_all(class_="o-story-content__main a-wysiwyg")[0].find_all("p")

    # Save article as docx
    set_docx(article_title, article_date, article_intro, [
        para.text for para in article_body], filename)